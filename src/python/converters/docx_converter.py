from .base import BaseConverter
from docx import Document
import re

class DocxConverter(BaseConverter):
    def get_supported_extensions(self) -> list[str]:
        return ['.docx', '.doc']

    def to_markdown(self, file_path: str, output_dir: str) -> str:
        doc = Document(file_path)
        markdown_lines = []

        for para in doc.paragraphs:
            style_name = para.style.name.lower() if para.style else ''
            text = para.text.strip()

            if not text:
                continue

            # 标题处理
            if 'heading' in style_name:
                level = int(style_name.replace('heading ', '')) if 'heading ' in style_name else 1
                markdown_lines.append(f"{'#' * level} {text}")
            # 列表处理
            elif para.style and 'list' in style_name:
                markdown_lines.append(f"- {text}")
            else:
                markdown_lines.append(text)

        # 处理表格
        for table in doc.tables:
            markdown_lines.append(self._table_to_markdown(table))

        return '\n\n'.join(markdown_lines)

    def _table_to_markdown(self, table) -> str:
        rows = []
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            rows.append('| ' + ' | '.join(cells) + ' |')
            if i == 0:
                rows.append('| ' + ' | '.join(['---'] * len(cells)) + ' |')
        return '\n'.join(rows)