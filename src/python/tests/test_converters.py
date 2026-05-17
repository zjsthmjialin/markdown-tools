import pytest
import os
import sys

# Add parent directory to path for imports
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from converters.markitdown_converter import MarkItDownConverter
from converters.pdf_converter import PDFConverter
from markdown.to_document import MarkdownToDocument

TEST_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
TEST_FILES = os.path.join(TEST_DIR, 'test_files')
OUTPUT_DIR = os.path.join(TEST_DIR, 'test_output')


@pytest.fixture(autouse=True)
def setup_output_dir():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    yield
    # Cleanup handled manually


def _test_file(name):
    path = os.path.join(TEST_FILES, name)
    if not os.path.exists(path):
        pytest.skip(f'Test file not found: {path}')
    return path


class TestMarkItDownConverter:
    converter = MarkItDownConverter()

    def test_docx_to_markdown(self):
        path = _test_file('sample.docx')
        result = self.converter.to_markdown(path, OUTPUT_DIR)
        assert isinstance(result, str)
        assert len(result) > 0

    def test_xlsx_to_markdown(self):
        path = _test_file('sample.xlsx')
        result = self.converter.to_markdown(path, OUTPUT_DIR)
        assert isinstance(result, str)
        assert len(result) > 0

    def test_pptx_to_markdown(self):
        path = _test_file('sample.pptx')
        result = self.converter.to_markdown(path, OUTPUT_DIR)
        assert isinstance(result, str)
        assert len(result) > 0

    def test_html_to_markdown(self):
        path = _test_file('sample.html')
        result = self.converter.to_markdown(path, OUTPUT_DIR)
        assert isinstance(result, str)
        assert len(result) > 0

    def test_txt_to_markdown(self):
        path = _test_file('sample.md')
        result = self.converter.to_markdown(path, OUTPUT_DIR)
        assert isinstance(result, str)
        assert len(result) > 0


class TestPDFConverter:
    converter = PDFConverter()

    def test_pdf_to_markdown(self):
        path = _test_file('sample.pdf')
        result = self.converter.to_markdown(path, OUTPUT_DIR)
        assert isinstance(result, str)
        assert len(result) > 0

    def test_pdf_supported_extensions(self):
        assert self.converter.get_supported_extensions() == ['.pdf']


class TestMarkdownToDocument:
    converter = MarkdownToDocument()
    sample_md = """# Test Document

## Section 1

This is a paragraph with **bold** and *italic* text.

- Item 1
- Item 2
- Item 3

## Section 2

| Name  | Age | City     |
|-------|-----|----------|
| Alice | 30  | Beijing  |
| Bob   | 25  | Shanghai |

```
def hello():
    print("Hello, World!")
```
"""

    def test_to_docx_with_tables(self):
        output = os.path.join(OUTPUT_DIR, 'test_tables.docx')
        self.converter.to_docx(self.sample_md, output)
        assert os.path.exists(output)
        assert os.path.getsize(output) > 0

        # Verify the document can be opened and has content
        from docx import Document
        doc = Document(output)
        assert len(doc.paragraphs) > 0

    def test_to_pptx(self):
        output = os.path.join(OUTPUT_DIR, 'test_output.pptx')
        self.converter.to_pptx(self.sample_md, output)
        assert os.path.exists(output)
        assert os.path.getsize(output) > 0

    def test_to_html(self):
        output = os.path.join(OUTPUT_DIR, 'test_output.html')
        self.converter.to_html(self.sample_md, output)
        assert os.path.exists(output)
        content = open(output, encoding='utf-8').read()
        assert '<table>' in content or '<th>' in content

    def test_to_pdf(self):
        output = os.path.join(OUTPUT_DIR, 'test_output.pdf')
        self.converter.to_pdf(self.sample_md, output)
        assert os.path.exists(output)
        assert os.path.getsize(output) > 0
